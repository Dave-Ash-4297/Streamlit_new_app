import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime
import re
import zipfile
import logging
import html

# --- Setup Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Constants ---
INDENT_FOR_IND_TAG_CM = 1.25
MAIN_LIST_TEXT_START_CM = 0.7
MARKER_OFFSET_CM = 0.7
SUB_LIST_TEXT_START_CM = 1.4
SUB_ROMAN_TEXT_START_CM = 2.1

# --- Utility Functions ---
def sanitize_input(text):
    """Escapes HTML characters in user input to prevent issues."""
    if not isinstance(text, str):
        text = str(text)
    return html.escape(text)

def add_formatted_runs(paragraph, text_line, placeholder_map):
    """
    Adds text to a paragraph, handling placeholders, bold, and underline tags.
    Handles <bd>, <ins>, and <***> (as bold).
    """
    processed_text = text_line
    for placeholder, value in placeholder_map.items():
        processed_text = processed_text.replace(f"{{{placeholder}}}", str(value))

    # Updated regex to include <***> tags
    parts = re.split(r'(<bd>|</bd>|<ins>|</ins>|<***>|</***>)', processed_text)
    is_bold = is_underline = False
    
    for part in parts:
        if not part: continue
        
        if part == "<bd>" or part == "<***>": 
            is_bold = True
        elif part == "</bd>" or part == "</***>": 
            is_bold = False
        elif part == "<ins>": 
            is_underline = True
        elif part == "</ins>": 
            is_underline = False
        else:
            run = paragraph.add_run(part)
            run.bold, run.underline = is_bold, is_underline
            run.font.name = 'Arial'
            run.font.size = Pt(11)

# --- Data Loading ---
@st.cache_data
def load_firm_details():
    """Loads static firm details."""
    return {
        "name": "Ramsdens Solicitors LLP", "short_name": "Ramsdens",
        "person_responsible_name": "Paul Pinder", "person_responsible_title": "Senior Associate",
        "supervisor_name": "Nick Armitage", "supervisor_title": "Partner",
        "person_responsible_phone": "01484 821558", "person_responsible_mobile": "07923 250815",
        "person_responsible_email": "paul.pinder@ramsdens.co.uk", "assistant_name": "Reece Collier",
        "supervisor_contact_for_complaints": "Nick Armitage on 01484 507121", "bank_name": "Barclays Bank PLC",
        "bank_address": "17 Market Place, Huddersfield", "account_name": "Ramsdens Solicitors LLP Client Account",
        "sort_code": "20-43-12", "account_number": "03909026",
        "marketing_email": "dataprotection@ramsdens.co.uk",
        "marketing_address": "Ramsdens Solicitors LLP, Oakley House, 1 Hungerford Road, Edgerton, Huddersfield, HD3 3AL"
    }

@st.cache_data
def load_precedent_text():
    """
    Loads the precedent text as a string.
    This is more reliable than reading from a file.
    """
    return """Our Ref: {matter_number}
Your Ref: {your_ref}
Date: {letter_date}
{client_name_input}
{client_address_line1}
{client_address_line2_conditional}
{client_postcode}

Dear {client_name_input},
<ins>What we have discussed and how we will charge for our work</ins>
1. Further to our recent discussions, we now write to confirm the terms under which Ramsdens Solicitors LLP <bd>(“Ramsdens”)</bd> will work for you in relation to this matter. As a firm that is regulated by the Solicitors Regulation Authority, we are required to send you this letter which contains specific and prescribed information. 
1. We enclose with this letter our Terms and Conditions of Business which must be read in conjunction with this letter. These documents are a formal communication and the language used is reflective of that. We hope that you understand. Please take the time to read these documents carefully. Where there is any conflict between this letter and our Terms and Conditions of Business, the terms of this letter will prevail. Your continuing instructions in this matter will amount to your acceptance of our Terms and Conditions of Business. 
<ins>Your Instructions</ins> 
1. We are instructed in relation to {qu1_dispute_nature} <bd>(“the Dispute”)</bd>. Per our recent discussions {qu2_initial_steps} <bd>(“the Work”)</bd>. 
1. This matter may develop over time and the nature of disputes is that opposing parties often seek to present facts and matters in a way that is favourable to their own case. We therefore cannot predict every eventuality but we will work with you to advise on any significant developments and review the overall strategy should that be required. Insofar as changes in the future may have a material impact on any cost estimates provided, we will discuss that with you. We will notify you of any material changes by telephone or in correspondence and we will of course always confirm any verbal advice in writing whenever you request that from us. 
<ins>Timescales</ins> 
1. {qu3_timescales}.
<ins>Action Required To Be Taken By You</ins>
<ins>Client Identification and Money Laundering</ins>
1. Solicitors are required by law to obtain evidence of their client’s identity and address to satisfy money laundering and client identification regulations. This includes clients that are corporate entities.
[corp]
1. We will make our own enquiries and obtain documentation from Companies House to identify our corporate client. If you believe Companies House’s records to be out of date, please let us know as soon as possible. We may also require documentation or information from the company itself.
1. We are also obliged to identify the individuals at the corporate client who provide us with instructions, which usually means directors of limited companies or members/partners in LLPs, and sometimes we must do the same for “beneficial owners”. Either situation may include you as the recipient of this letter, but may also include other people at the business. We will tell you who else may be required to provide identification.
[/corp]
[indiv]
1. To comply with the individual identity requirement, you have two options:
<a>We can carry out a remote ID verification of you and your ID documents using a SmartSearch facility. If you would like us to verify your identification remotely please provide your name, address, date of birth, personal email and mobile number. Once the search has been undertaken, SmartSearch will send you a text or email with a link to use on your smartphone which will require you to take a photo of your ID document and then yourself which it will then upload to its system that will check the document, provide us with a copy, and verify that you are the person on the ID document. Or;
<a>You can provide us with two documents referred to in the list below, one photographic and one showing your current address. If you are local to any of our offices please call with your original documents and they will be copied whilst you wait and the copies forwarded to us:
<i>Current signed passport;
<i>Household utility bill;
<i>Residence permit issued by the Home Office to EEA nationals on sight of own country passport;
<i>Current UK or EEA photo-card driving licence; or
<i>National Identity Card containing a photograph.
[/indiv]
1. Please note that until these identification requirements have been satisfied, we may not be ableto accept any money from you or make any substantial progress with your matter. It is therefore important that you provide your documents as soon as possible to avoid any delays.
<ins>Document Preservation and Disclosure</ins>
1. In the event that your matter is litigated before a Court, all parties will be required to give full disclosure of all material relevant to the Dispute. It is therefore essential that you preserve any and all such material that includes correspondence, documents, emails, text and SMS messages and/or other electronic communications of any sort. Your disclosure obligations include an obligation to disclose material that may harm your case or help your opponent’s case, as well as those on which you may rely or which help. If any device on which any such material is stored is to be disposed of or ceases to be used, you must ensure that copies are kept of the material.
<ins>People Responsible For Your Case</ins>
1. I shall be the person with responsibility for your case. My name is {person_responsible_name} and I am a {person_responsible_title} with the firm. My work will be carried out under the supervision of {supervisor_name} who is a {supervisor_title} of the firm.
1. The easiest way to communicate with me will be either by telephone on {person_responsible_phone}, my mobile {person_responsible_mobile}, or via email to {person_responsible_email}.
1. There may be occasions when I am not immediately available to speak or meet with you and in these circumstances you should ask to speak to my Assistant, {assistant_name} who will be able to help you.
1. At Ramsdens we aim to provide the best possible service to our clients and in order to do this we may arrange for one of our client care team to contact you to discuss how we are doing and what we might do better. Please let us know if you would prefer not to be contacted by our team during our handling of your matter. We do however, need to know from you if you feel dissatisfied about the service you are receiving. Should you have any occasion to feel unhappy about our service please let me know straight away and I will discuss this with you. If you are unable to resolve matters with me and still have concerns regarding our service, contact {supervisor_contact_for_complaints} who will attempt to resolve your concerns with you. Formal complaints will be dealt with in accordance with our Firm's complaints procedures which can be provided on request. In the event you are not satisfied with our handling of your complaint you can contact the Legal Ombudsman, full details will be given as part of our complaints procedure.
1. You also have a right to complain about any bill sent by us by applying to the Court for an assessment of the bill under Part III of the Solicitors Act 1974.
<ins>Costs and Disbursements</ins>
<ins>Costs</ins>
1. Our charges to you will be calculated and incurred on a time-spent basis. Time will be recorded on your matter in units of six minutes for letters (generally representing a unit per page or part thereof), emails written (again, representing a unit per equivalent to a page of normal letter) and telephone calls made and received.
1. Our current hourly charge-out rates, exclusive of VAT, are as follows:
[FEE_TABLE_PLACEHOLDER]
1. Our hourly charge-out rates are reviewed periodically and we will notify you of any increases. We will also notify you of any changes in the status of legal personnel and their hourly charge-out rate. Unless otherwise agreed with you, we will account to you every month for the fees that have been incurred in relation to this matter. If you require an up to date statement of fees incurred at any time then please ask us and we will provide you with that information. Unless otherwise stated, interim bills are on account of costs and are usually prepared taking into account the value of the time recorded against the matter as at the date of the interim bill. If we hold any monies on account of your costs when an invoice is raised, these monies will be utilised towards discharging the invoice.
<ins>Disbursements</ins>
1. Our hourly charge-out rates do not include expenses for which we will be responsible on your behalf. These expenses are referred to as disbursements and may include travel or accommodation expenses, Court fees or the costs of Barristers or expert witnesses. Where possible, we will endeavour to seek your permission prior to instructing a third party in relation to your matter.
1. We will not pay out any disbursements on your behalf until the monies have been paid by you.
<ins>Legal Expenses Insurance</ins>
[indiv]
1. It may be that you or a member of your household has the benefit of legal expenses insurance that might cover you for legal costs in connection with this matter. If you wish us to check your eligibility, please let us have a copy of the relevant insurance schedule and policy document. Alternatively you may be entitled to have your liability for costs paid by another person; for example, an employer or Trade Union. Again, please let us know if you wish us to assist you in checking such eligibility. Please note that you will remain responsible for our charges until such time as any legal expenses insurers have agreed to cover you for our legal costs.
[/indiv]
[corp]
1. It may be possible to purchase “After the Event” legal expenses insurance cover to cover your opponents, or, possibly, your costs in this matter. If you wish to explore the possibility of such insurance cover, please let us know. Please bear in mind, though, that there will be costs involved in making an application for cover, and it is likely that a large premium (the amount of which will depend on the amount of costs protected and the prospects of success) will be payable at the outset and possibly on any subsequent anniversary of the inception of the policy.
[/corp]
<ins>Your Costs Responsibility to Ramsdens</ins>
1. Our charges to you are not contingent upon the result of your case. You are primarily responsible for the payment of our costs and disbursements. Whilst we may be able to recover a portion of your costs from your opponent, this is not always possible and does not affect your primary responsibility to pay our costs and disbursements.
<ins>Section 74 Solicitors Act 1974 Agreement & Recovery of Costs</ins>
1. It is common in litigation that even where costs are recoverable from an opponent, such recovery will not equate to the level of costs incurred by the successful party. Our agreement expressly permits us to charge an amount of costs greater than that which you will recover or could have recovered from your opponent, and expressly permits payment of such sum.
1. This part of our agreement is made under section 74(3) of the Solicitors Act 1974 and Civil Procedure Rules 46.9 (2) and (3).
1. If a Court orders your opponent to pay your costs, you should be aware that:
<a>You will have to pay the costs to us in the first instance and thereafter, if you are successful in your claim, seek to recover those costs from the other side;
<a>You are unlikely to recover the entirety of our charges from the other side even if your claim is successful;
<a>In the unlikely event that your claim is subject to the fixed recoverable costs regime (see below) and the fixed costs recoverable from your opponent exceed the level of our charges calculated and incurred on a time-spent basis, you agree that the charges due to us from you will be the amount of fixed costs recoverable from your opponent.
<a>Your opponent may refuse to comply with the Court’s order. If they do not pay, then you may seek to enforce the Court’s order (for example by sending in the bailiffs or obtaining a charge over property owned by them). However, you should be aware that this itself costs more money and takes time.
<a>Your opponent may have very little by way of assets or they may simply disappear. If this happens then you will not be able to recover your costs or indeed any other monies awarded to you. That is why it is important that in financial disputes you consider now whether your opponent has sufficient assets to pay you a lump sum or instalments as appropriate.
<a>There may be points during your case (including at its conclusion) where you are successful only in part on the issues in it, as a result of which you are entitled to payment of some of your costs by your opponent.
<a>If your opponent receives funding from the Community LegalService, there are statutory controls on the amount of costs that can be recovered from them. In these circumstances, it is highly unlikely that the Court will make an order that your opponent would have to contribute anything to your costs.
<ins>Fixed Recoverable Costs</ins>
1. Depending upon the value and complexity of a claim, the Court will allocate it to one of four ‘tracks’ when managing the case. If a claim is successful and a Court orders one party to pay the other’s costs, the amount of the costs that can be recovered may be fixed by the Court.
[a1]
1. From the information that you have supplied us with, the claim has already been allocated to the Small Claims Track which is the normal track for claims with a monetary value of £10,000 or less. Having been allocated to the Small Claims Track, the normal rule is that only the following limited costs are recoverable by a successful party:
<a>Any Court fees paid.
<a>Fixed issue costs ranging between £50 and £125.
<a>Loss of earnings not exceeding £95 per person per day.
<a>Expenses reasonably incurred in travelling to and from and attending a Court hearing.
<a>A sum not exceeding £750 for any expert’s fees.
1. There are some exceptions to the normal rule and the Court can award costs against a party that has acted unreasonably. However, in practice such awards are rare.
[/a1]
[a2]
1. From the information that you have supplied us with, the claim has already been allocated to the Fast Track which is the normal track for claims with a monetary value of between £10,000 and £25,000. Having been allocated to the Fast Track, the Court has also assigned your/your opponent’s claim to a Band 1/2/3/4. This means that as the Claimant/Defendant in the proceedings PAUL REVIEW HERE.
[/a2]
[a3]
1. From the information that you have supplied us with, the claim has already been allocated to the Intermediate Track which is the normal track for claims with a monetary value of between £25,000 and £100,000. Having been allocated to the Intermediate Track, the Court has also assigned your/your opponent’s claim to Band 1/2/3/4. This means that as the Claimant/Defendant in the proceedings, we know that the costs that may be recoverable from your opponent/you will be fixed dependent upon the stage of the proceedings in which the claim is resolved. A table setting out these fixed recoverable costs is enclosed with this letter.
[/a3]
[a4]
1. From the information that you have supplied us with, the claim has already been allocated to the Multi-Track which is the normal track for claims with a monetary value of over £100,000. Having been allocated to the Multi-Track, this means that the fixed costs regime does not apply to your/your opponent’s claim and the general rule that the ‘loser pays the winner’s costs’ will apply, subject to any costs budgeting that has been implemented by the Court and the caveats set out above under the heading <***>Section 74 Solicitors Act 1974 Agreement & Recovery of Costs</***>.
[/a4]
[u1]
1. From the information that you have supplied us with, it is likely that were Court proceedings to be commenced, the claim would be allocated to the Small Claims Track which is the normal track for claims with a monetary value of £10,00D00 or less. Upon allocation to the Small Claims Track, the normal rule is that only the following limited costs are recoverable by a successful party:
<a>Any Court fees paid.
<a>Fixed issue costs ranging between £50 and £125.
<a>Loss of earnings not exceeding £95 per person per day.
<a>Expenses reasonably incurred in travelling to and from and attending a Court hearing.
<a>A sum not exceeding £750 for any expert’s fees.
1. There are some exceptions to the normal rule and the Court can award costs against a party that has acted unreasonably. However, in practice such awards are rare.
[/u1]
[u2]
1. From the information that you have supplied us with, it is likely that were Court proceedings to be commenced, the claim would be allocated to the Fast Track which is the normal track for claims with a monetary value of between £10,000 and £25,000. Upon allocation to the Fast Track, the Court will assign your/your opponent’s claim to one of four ‘bands’ depending upon the complexity and number of issues in the claim. When the claim is assigned, as the Claimant/Defendant in the proceedings, we know that the costs that may be recoverable from your opponent/you will be fixed dependent upon the stage of the proceedings in which the claim is resolved. A table setting out these fixed recoverable costs is enclosed with this letter.
[/u2]
[u3]
1. From the information that you have supplied us with, it is likely that were Court proceedings to be commenced, the claim would be allocated to the Intermediate Track which is the normal track for claims with a monetary value of between £25,000 and £100,000. Upon allocation to the Intermediate Track, the Court will assign your/your opponent’s claim to one of four ‘bands’ depending upon the complexity and number of issues in the claim. When the claim is assigned, as the Claimant/Defendant in the proceedings, we know that the costs that may be recoverable from your opponent/you will be fixed dependent upon the stage of the proceedings in which the claim is resolved. A table setting out these fixed recoverable costs is enclosed with this letter.
[/u3]
[u4]
1. From the information that you have supplied us with, it is likely that were Court proceedings to be commenced, the claim would be allocated to the Multi-Track which is the normal track for claims with a monetary value of in excess of £100,000. Upon allocation to the Multi-Track, the fixed costs regime will not apply to your/your opponent’s claim and the general rule that the ‘loser pays the winner’s costs’ will apply, subject to any costs budgeting that has been implemented by the Court and the caveats set out above under the heading <***>Section 74 Solicitors Act 1974 Agreement & Recovery of Costs</***>.
[/u4]
<ins>Costs Advice</ins>
1. From the information you have provided us with to date, we estimate that our costs for the initial stage of the Work will be {qu4_initial_costs_estimate}. If any further work is required thereafter, we will discuss the likely associated costs with you beforehand.
1. It is always difficult to give an indication of the likely costs to be incurred in cases of this type. This is because it is impossible to say at this stage when the case may be brought to a conclusion and the amount of work that may be required to reach that point. The vast majority of cases are settled without the need for Court proceedings and of those where Court proceedings are commenced, the majority are settled without a trial. The actual amount of costs to be incurred will depend upon the arguments being advanced and the amount and nature of the evidence involved. The more evidence that is required, the greater the amount of time that will be spent on it by the parties and the Court and, therefore, the greater the costs.
1. The involvement of expert evidence (such as in the form of valuation evidence) will also contribute to an increase in the costs involved.
1. In the event that it may appear that our initial estimate of costs may be exceeded, we will notify you of these changes. We will review our estimate of costs at least every six months.
1. There may be occasions during the conduct of your case where significant disbursements or major amounts of chargeable time are due to be incurred. We reserve the right to seek payment in advance for these commitments, and routinely do so. In the event that we do seek such payment in advance and it is not made by any reasonable deadline set, we reserve the right to cease acting for you in this matter. In the event that we do cease to act we would attempt to mitigate the impact that doing so would have on your case but it is possible that your case may be prejudiced as a Zresult. We also reserve the right to cease acting for you in the event that any bills rendered to you are not paid within the timescale required.
1. To this extent, you agree with us that our retainer in this matter is not to be considered an entire agreement, such that we are not obliged to continue acting for you to the conclusion of the matter and are entitled to terminate your retainer before your case is concluded. We are required to make this clear because there has been legal authority that in the absence of such clarity a firm was required to continue acting in a case where they were no longer being funded to do so.
1. You have a right to ask for your overall cost to be limited to a maximum and we trust you will liaise with us if you wish to limit your costs. We will not then exceed this limit without first obtaining your consent. However this does mean that your case may not be concluded if we have reached your cost limit.
1. In Court or some Tribunal proceedings, you may be ordered to pay the costs of someone else, either in relation to the whole of the costs of the case if you are unsuccessful or in relation to some part or issue in the case. Also, you may be ordered to pay the costs of another party during the course of proceedings in relation to a particular application to the Court. In such case you will need to provide this firm with funds to discharge such liability within seven days as failure to do so may prevent your case progressing. Please be aware that once we issue a Court or certain Tribunal claims or counterclaim on your behalf, you are generally unable to discontinue your claim or counterclaim without paying the costs of your opponent unless an agreement on costs is reached.
<ins>Limitation of Liability</ins>
1. The liability of Ramsdens Solicitors LLP, its partners and employees in any circumstances whatsoever, whether in contract, tort, statute or otherwise and howsoever caused (including our negligence) for loss or damage arising from or in connection with the provision of services to you shall be limited to the sum of £3,000,000.00 (three million pounds) excluding costs and interest.
<ins>Bank Accounts and Cybercrime</ins>
1. Should we ask you to pay money to us during the course of your matter then please send your funds to our account held with {bank_name} at {bank_address} to:
[ind]Account Name: {account_name}
[ind]Sort Code: {sort_code}
[ind]Account Number: {account_number}
1. Should you receive any email correspondence regarding our bank account details please telephone your usual contact at Ramsdens before sending your first payment to verify that the details you have been given are correct. We would never advise our clients of any change in our bank account details by email. Should this happen please treat the email as suspicious and contact us immediately. Please do not send any funds until you have verified that the details are correct.
1. Similarly, if an occasion arises whereby we need to send money to you, we will not accept your bank account details by email without further verification. It is likely that we will telephone you to confirm that the details supplied to us are correct.
<ins>Quality Standard</ins>
1. Our firm is registered under the Lexcel quality standard of the Law Society. As a result of this we are or may become subject to periodic checks by outside assessors. This could mean that your file is selected for checking, in which case we would need your consent for inspection to occur. All inspections are, of course, conducted in confidence. If you prefer to withhold consent, work on your file will not be affected in any way. Since very few of our clients do object to this we propose to assume that we do have your consent unless you notify us to the contrary. We will also assume, unless you indicate otherwise, that consent on this occasion will extend to all future matters which we conduct on your behalf. Please do not hesitate to contact us if we can explain this further or if you would like us to mark your file as not to be inspected. Alternatively if you would prefer to withhold consent please put a line through this section in the copy letter and return to us.
<ins>Data Protection</ins>
1. The enclosed Privacy Notice explains how and why we collect, store, use and share your personal data. It also explains your rights in relation to your personal data and how to contact us or supervisory authorities in the event you have a complaint. Please read it carefully. This Privacy Notice is also available on our website, www.ramsdens.co.uk.
1. Our use of your personal data is subject to your instructions, the EU General Data Protection Regulation (GDPR), other relevant UK and EU legislation and our professional duty of confidentiality. Under data-protection law, we can only use your personal data if we have a proper reason for doing so. Detailed reasons why we may process your personal data are set out in our Privacy Notice but examples are:
<a>To comply with our legal and regulatory obligations;
<a>For the performance of our contract with you or to take steps at your request before entering into a contract; or
<a>For our legitimate interests or those of a third party, including:
<i>Operational reasons, such as recording transactions, training and quality control;
<i>Updating and enhancing client records;
<i>Analysis to help us manage our practice; and
<i>Marketing, such as by sending you updates about subjects and/or events that may be of interest to you.
1. However, this does not apply to processing sensitive personal data about you, as defined. If it is necessary to process this data for the continued provision of our services to you, we will need your explicit consent for doing so and will request this from you as required.
<ins>Marketing Communications</ins>
1. We would like to use your personal data to send you updates (by email, telephone or post) about legal developments that might be of interest to you and/or information about our services.
1. This will be done pursuant to our Privacy Notice (referred to above), which contains more information about our and your rights in this respect.
1. You have the right to opt out of receiving promotional communications at any time, by:
<a>Contacting us by email on {marketing_email};
<a>Using the ‘unsubscribe’ link in emails; or
<a>Writing to Marketing Department at: {marketing_address}.

Yours sincerely,


{name}
Solicitor
"""

# --- Document Generation Logic ---
def generate_client_care_document(precedent_content, app_inputs):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    numbering_elm = doc.part.numbering_part.element
    abstract_num_id, num_instance_id = 10, 1

    def setup_numbering_style(numbering_element):
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))

        def create_level(ilvl, numFmt, lvlText, left_cm):
            lvl = OxmlElement('w:lvl')
            lvl.set(qn('w:ilvl'), str(ilvl))
            lvl.append(OxmlElement('w:start', {qn('w:val'): '1'}))
            lvl.append(OxmlElement('w:numFmt', {qn('w:val'): numFmt}))
            lvl.append(OxmlElement('w:lvlText', {qn('w:val'): lvlText}))
            pPr = OxmlElement('w:pPr')
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(Cm(left_cm).twips))
            ind.set(qn('w:hanging'), str(Cm(MARKER_OFFSET_CM).twips))
            pPr.append(ind)
            lvl.append(pPr)
            return lvl

        abstract_num.append(create_level(0, 'decimal', '%1.', MAIN_LIST_TEXT_START_CM))
        abstract_num.append(create_level(1, 'lowerLetter', '%2.', SUB_LIST_TEXT_START_CM))
        abstract_num.append(create_level(2, 'lowerRoman', '%3.', SUB_ROMAN_TEXT_START_CM))
        numbering_element.append(abstract_num)

        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_instance_id))
        num.append(OxmlElement('w:abstractNumId', {qn('w:val'): str(abstract_num_id)}))
        numbering_element.append(num)

    setup_numbering_style(numbering_elm)
    placeholder_map = app_inputs['placeholder_map']

    def add_list_item(text, level):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numPr.get_or_add_ilvl().val = level
        numPr.get_or_add_numId().val = num_instance_id
        add_formatted_runs(p, text, placeholder_map)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    current_block_tag = None
    lines = precedent_content.splitlines()
    for line in lines:
        stripped_line = line.strip()
        match_start_tag = re.match(r'\[(indiv|corp|a[1-4]|u[1-4])\]', stripped_line)
        match_end_tag = re.match(r'\[/(indiv|corp|a[1-4]|u[1-4])\]', stripped_line)

        if match_start_tag:
            current_block_tag = match_start_tag.group(1)
            continue
        if match_end_tag:
            current_block_tag = None
            continue
        
        should_render = True
        if current_block_tag:
            tag = current_block_tag
            claim_assigned = app_inputs['claim_assigned']
            selected_track = app_inputs['selected_track']
            tag_map = {'a1': (True, "Small Claims Track"), 'a2': (True, "Fast Track"), 'a3': (True, "Intermediate Track"), 'a4': (True, "Multi Track"), 'u1': (False, "Small Claims Track"), 'u2': (False, "Fast Track"), 'u3': (False, "Intermediate Track"), 'u4': (False, "Multi Track")}
            if tag in ['indiv', 'corp']:
                should_render = (tag == 'indiv' and app_inputs['client_type'] == 'Individual') or (tag == 'corp' and app_inputs['client_type'] == 'Corporate')
            elif tag in tag_map:
                expected_assignment, expected_track = tag_map[tag]
                should_render = (claim_assigned == expected_assignment and selected_track == expected_track)

        # Skip empty lines (but not if they are part of a block being skipped)
        if not stripped_line and current_block_tag is None:
            doc.add_paragraph() # Add blank lines as they appear in the template
            continue

        if not should_render or not stripped_line:
            continue
        
        match_heading = re.match(r'^<ins>(.*)</ins>$', stripped_line)
        match_numbered_list = re.match(r'^(\d+)\.\s*(.*)', stripped_line)
        match_letter_list = re.match(r'^<a>\s*(.*)', stripped_line)
        match_roman_list = re.match(r'^<i>\s*(.*)', stripped_line)

        if stripped_line == '[FEE_TABLE_PLACEHOLDER]':
            for fee_line in app_inputs['fee_table']:
                add_list_item(fee_line, level=0)
        elif match_heading:
            p = doc.add_paragraph()
            add_formatted_runs(p, f"<ins>{match_heading.group(1)}</ins>", placeholder_map)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif match_numbered_list:
            add_list_item(match_numbered_list.group(2), level=0)
        elif match_letter_list:
            add_list_item(match_letter_list.group(1), level=1)
        elif match_roman_list:
            add_list_item(match_roman_list.group(1), level=2)
        else:
            p = doc.add_paragraph()
            # Handle [ind] tag for indentation, but render the rest of the line
            cleaned_content = line.replace('[ind]', '').lstrip() # Use lstrip to remove leading spaces
            if '[ind]' in line:
                p.paragraph_format.left_indent = Cm(INDENT_FOR_IND_TAG_CM)
            
            add_formatted_runs(p, cleaned_content, placeholder_map)
            
            # Smart default alignment
            if '{' not in line and '}' not in line and not re.search(r'<.*?>', line):
                # Apply justify if it's likely a simple text paragraph
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Smart spacing
            if line.startswith("Dear") or line.startswith("Yours sincerely"):
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith("{name}") or line.startswith("Solicitor"):
                 p.paragraph_format.space_after = Pt(0)
            else:
                p.paragraph_format.space_after = Pt(12)
                
    return doc

def generate_initial_advice_doc(app_inputs):
    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Arial', Pt(11)
    p = doc.add_paragraph()
    add_formatted_runs(p, "Initial Advice Summary - Matter Number: {matter_number}", app_inputs['placeholder_map'])
    p.paragraph_format.space_after = Pt(12)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    advice_date = app_inputs['initial_advice_date'].strftime('%d/%m/%Y') if app_inputs.get('initial_advice_date') else ''
    rows_data = [("Date of Advice", advice_date), ("Method of Advice", app_inputs.get('initial_advice_method', '')), ("Advice Given", app_inputs.get('initial_advice_content', ''))]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- Streamlit UI ---
st.set_page_config(layout="wide", page_title="Ramsdens Client Care Letter Generator")
st.title("Ramsdens Client Care Letter Generator")

firm_details = load_firm_details()
precedent_content = load_precedent_text()
if not precedent_content:
    st.error("Failed to load precedent text. App cannot continue.")
    st.stop()

# --- Initialize session state for cost hours ---
if 'lower_hours' not in st.session_state:
    st.session_state.lower_hours = 2.0  # Starts at 2x hourly rate
if 'upper_hours' not in st.session_state:
    st.session_state.upper_hours = 3.5  # Starts at 3.5x hourly rate
if 'fixed_hours' not in st.session_state:
    st.session_state.fixed_hours = 2.5

# --- Callback functions for buttons ---
def increment(key):
    st.session_state[key] += 0.5

def decrement(key):
    if st.session_state[key] > 0.5: # Prevent going to 0 or negative
        st.session_state[key] -= 0.5

# --- Interactive Cost Estimation Section (Placed Before the Form) ---
st.header("Cost Estimation")
st.write("Adjust the estimated hours for the initial work. The cost will update automatically.")
hourly_rate = st.number_input("Your Hourly Rate (£)", value=295, step=5, key="hourly_rate_input")
cost_type_is_range = st.toggle("Use a cost range", True)

if cost_type_is_range:
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("##### Lower Estimate")
        b1, b2, b3 = st.columns([1, 2, 1])
        with b1:
            st.button("➖", key="dec_lower", on_click=decrement, args=('lower_hours',), use_container_width=True)
        with b2:
            cost = st.session_state.lower_hours * hourly_rate
            st.metric(label="Hours", value=f"{st.session_state.lower_hours:.1f}", delta=f"£{cost:,.2f} + VAT")
        with b3:
            st.button("➕", key="inc_lower", on_click=increment, args=('lower_hours',), use_container_width=True)
    with c2:
        st.markdown("##### Upper Estimate")
        b1, b2, b3 = st.columns([1, 2, 1])
        with b1:
            st.button("➖", key="dec_upper", on_click=decrement, args=('upper_hours',), use_container_width=True)
        with b2:
            cost = st.session_state.upper_hours * hourly_rate
            st.metric(label="Hours", value=f"{st.session_state.upper_hours:.1f}", delta=f"£{cost:,.2f} + VAT")
        with b3:
            st.button("➕", key="inc_upper", on_click=increment, args=('upper_hours',), use_container_width=True)
else:
    st.markdown("##### Fixed Fee Estimate")
    _, c, _ = st.columns([1, 2, 1])
    with c:
        b1, b2, b3 = st.columns([1, 2, 1])
        with b1:
            st.button("➖", key="dec_fixed", on_click=decrement, args=('fixed_hours',), use_container_width=True)
        with b2:
            cost = st.session_state.fixed_hours * hourly_rate
            st.metric(label="Hours", value=f"{st.session_state.fixed_hours:.1f}", delta=f"£{cost:,.2f} + VAT")
        with b3:
            st.button("➕", key="inc_fixed", on_click=increment, args=('fixed_hours',), use_container_width=True)

# --- Data Input Form ---
with st.form("input_form"):
    st.header("1. Letter & Client Details")
    c1, c2 = st.columns(2)
    with c1:
        our_ref = st.text_input("Our Reference", "PDP/10011/001")
        your_ref = st.text_input("Your Reference", "REF")
        letter_date = st.date_input("Letter Date", datetime.today())
    with c2:
        client_name_input = st.text_input("Client Full Name / Company Name", "Mr. John Smith")
        client_address_line1 = st.text_input("Address Line 1", "123 Example Street")
        client_address_line2 = st.text_input("Address Line 2 (optional)", "SomeTown")
        client_postcode = st.text_input("Postcode", "EX4 MPL")
        client_type = st.radio("Client Type", ("Individual", "Corporate"), horizontal=True)

    st.header("2. Initial Advice & Case Details")
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("Initial Advice Summary")
        initial_advice_content = st.text_area("Advice Given", "Advised on merits...", height=100)
        initial_advice_method = st.selectbox("Method", ["Phone Call", "In Person", "Teams Call"])
        initial_advice_date = st.date_input("Date", datetime.today())
    with c2:
        st.subheader("Case Track")
        claim_assigned_input = st.radio("Is claim already assigned?", ("No", "Yes"), horizontal=True, index=1) # Default to Yes
        selected_track = st.selectbox("Which track applies?", ["Small Claims Track", "Fast Track", "Intermediate Track", "Multi Track"], index=3) # Default to Multi

    st.header("3. Dynamic Content")
    qu1_dispute_nature = st.text_area('Dispute Nature', "a contractual matter", height=75)
    qu2_initial_steps = st.text_area('Initial Work', "we agreed I would review documentation", height=100)
    qu3_timescales = st.text_area("Estimated Timescales", "The initial part of the Work will take around two to four weeks.", height=100)
    
    # Final submit button for the form
    submitted = st.form_submit_button("Generate Documents")

if submitted:
    # 1. Collate all inputs and generate final cost text
    if cost_type_is_range:
        lower_cost = st.session_state.lower_hours * hourly_rate
        upper_cost = st.session_state.upper_hours * hourly_rate
        costs_text = f"£{lower_cost:,.2f} to £{upper_cost:,.2f} plus VAT"
    else:
        fixed_cost = st.session_state.fixed_hours * hourly_rate
        costs_text = f"a fixed fee of £{fixed_cost:,.2f} plus VAT"

    roles = [("Partner", hourly_rate * 1.5), ("Senior Associate", hourly_rate), ("Associate", hourly_rate * 0.8), ("Trainee", hourly_rate * 0.5)]
    fee_table = [f"{role}: £{rate:,.2f} per hour (excl. VAT)" for role, rate in roles]

    app_inputs = {
        'client_type': client_type,
        'claim_assigned': claim_assigned_input == "Yes",
        'selected_track': selected_track,
        'fee_table': fee_table,
        'initial_advice_content': initial_advice_content,  
        'initial_advice_method': initial_advice_method,
        'initial_advice_date': initial_advice_date
    }
    
    placeholder_map = {
        'matter_number': sanitize_input(our_ref),
        'your_ref': sanitize_input(your_ref),
        'letter_date': letter_date.strftime('%d %B %Y'),
        'client_name_input': sanitize_input(client_name_input),
        'client_address_line1': sanitize_input(client_address_line1),
        'client_address_line2_conditional': sanitize_input(client_address_line2) if client_address_line2 else "",
        'client_postcode': sanitize_input(client_postcode),
        'qu1_dispute_nature': sanitize_input(qu1_dispute_nature),
        'qu2_initial_steps': sanitize_input(qu2_initial_steps),
        'qu3_timescales': sanitize_input(qu3_timescales),
        'qu4_initial_costs_estimate': costs_text,
        'name': sanitize_input(firm_details["person_responsible_name"]), # Special case for 'name'
    }
    # Add all other firm details to the map
    placeholder_map.update(firm_details)
    app_inputs['placeholder_map'] = placeholder_map

    # 2. Generate documents
    try:
        care_letter_doc = generate_client_care_document(precedent_content, app_inputs)
        client_care_doc_io = io.BytesIO()
        care_letter_doc.save(client_care_doc_io)
        client_care_doc_io.seek(0)
        
        advice_doc_io = generate_initial_advice_doc(app_inputs)
        
        client_name_safe = re.sub(r'[^\w\s-]', '', client_name_input).strip().replace(' ', '_')
        zip_io = io.BytesIO()
        with zipfile.ZipFile(zip_io, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.writestr(f"Client_Care_Letter_{client_name_safe}.docx", client_care_doc_io.getvalue())
            zipf.writestr(f"Initial_Advice_Summary_{client_name_safe}.docx", advice_doc_io.getvalue())
        zip_io.seek(0)
        
        # --- FIX: Store the zip file in session state ---
        st.session_state.zip_buffer = zip_io.getvalue()
        st.session_state.client_name_safe = client_name_safe
        
        st.success("✅ Documents Generated Successfully! Download button is below.")
        
        # --- REMOVED st.download_button from here ---

    except Exception as e:
        st.error(f"An error occurred: {e}")
        logger.exception("Error during document generation:")

# --- FIX: Add the download button here, outside the form logic ---
# This checks if a file has been generated and is waiting in session state.
if "zip_buffer" in st.session_state and st.session_state.zip_buffer is not None:
    st.download_button(
        label="Download All Documents as ZIP",
        data=st.session_state.zip_buffer,
        file_name=f"Client_Docs_{st.session_state.client_name_safe}.zip",
        mime="application/zip",
        # Clear the state after download so the button disappears
        on_click=lambda: st.session_state.pop("zip_buffer", None)
    )
